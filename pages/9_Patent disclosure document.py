import streamlit as st
import ollama as ol
import datetime
import json
import requests
import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import io
from docx.shared import Inches
import subprocess
from bs4 import BeautifulSoup
import time

def log_interaction(action, data):
    # 日志记录功能
    timestamp = datetime.datetime.now().isoformat()  # 获取当前时间的ISO格式
    log = {"timestamp": timestamp, "action": action, "data": data}  # 创建一个包含时间戳、动作和数据的字典
    with open("user_interactions_log.json", "a") as logfile:  # 以追加模式打开日志文件
        logfile.write(json.dumps(log) + "\n")  # 将日志记录写入文件，并添加换行符


def print_txt(text):
    # 定义一个函数，以正确方向显示阿拉伯文本
    if any("\u0600" <= c <= "\u06FF" for c in text):  # 检查文本中是否包含阿拉伯字符
        text = f"<p style='direction: rtl; text-align: right;'>{text}</p>"  # 如果包含，则设置文本为从右到左显示
    st.markdown(text, unsafe_allow_html=True)  # 使用Markdown显示文本，并允许HTML


def print_chat_message(message):
    # 定义打印聊天信息的函数
    text = message["content"]  # 获取信息内容
    if message["role"] == "user":  # 如果信息来自用户
        with st.chat_message("user", avatar="🎙️"):  # 设置用户头像并显示信息
            print_txt(text)  # 调用print_txt函数以处理文本方向
    elif message["role"] == "assistant":  # 如果信息来自助手
        with st.chat_message("assistant", avatar="🦙"):  # 设置助手头像并显示信息
            print_txt(text)  # 调用print_txt函数以处理文本方向


def init_chat_history(key, system_prompt):
    # 定义打印聊天信息的函数
    if "chat_history" not in st.session_state:  # 如果会话状态中没有聊天记录
        st.session_state.chat_history = {}  # 初始化聊天记录
    if key not in st.session_state.chat_history:  # 如果指定的聊天记录不存在
        st.session_state.chat_history[key] = [{"role": "system", "content": system_prompt}]  # 创建新的聊天记录并添加系统提示


def read_docx(file):
    # 读取DOCX文件
    doc = Document(file)  # 打开Word文档
    full_text = []  # 初始化一个空列表以存储文本
    for para in doc.paragraphs:  # 遍历文档中的所有段落
        full_text.append(para.text)  # 将段落文本添加到列表中
    return '\n'.join(full_text)  # 将所有段落文本合并为一个字符串并返回

def remove_symbols_from_word(doc_bytes):
    # 从内存中的字节创建 Word 文档
    buffer = io.BytesIO(doc_bytes)
    original_doc = Document(buffer)

    # 创建一个新的 Word 文档
    new_doc = Document()

    # 处理每个段落
    for paragraph in original_doc.paragraphs:
        # 去除无关符号
        clean_text = re.sub(r'[\*\#\-\[\]]', '', paragraph.text)
        # 将清理后的文本添加到新文档中
        new_doc.add_paragraph(clean_text)

    # 创建一个新的 BytesIO 对象来存储处理后的文档
    output_buffer = io.BytesIO()
    new_doc.save(output_buffer)
    output_buffer.seek(0)
    processed_doc_bytes = output_buffer.getvalue()

    return processed_doc_bytes

def search_patents(x1):
    post_url = 'https://www.chonghus.com/hxapi/qc/zytxt'
    headers = {
        'User-Agent': 'Mozilla/5.0',
        'Content-Type': 'application/json'
    }
    data = {
        'type': 'search_zy_txt',
        'Context': f"""{x1}""",
        'total': '10',
        'title': '摘要全文查重'
    }
    try:
        response = requests.post(post_url, headers=headers, json=data, timeout=10)
        time.sleep(1)
        response.raise_for_status()  # 检查响应是否成功
    except requests.exceptions.RequestException as e:
        print(f"请求失败: {e}")
        exit()
    if response.status_code == 200:
        soup = BeautifulSoup(response.content, 'lxml')
        json_text = soup.p.text
        data = json.loads(json_text)
        unique_patents = {}
        for index, item in enumerate(data['msg']):
            unique_patents[str(index)] = item
        temp_dict = {}
        for key, value in unique_patents.items():
            pat_name = value['pat_name']
            if pat_name not in temp_dict:
                temp_dict[pat_name] = value
        unique_patents = {str(i): v for i, v in enumerate(temp_dict.values())}
        for key in unique_patents:
            uuid = unique_patents[key]['uuid']
            url = f'https://www.chonghus.com/hxapi2/pat/detail?id={uuid}'
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36'
            }
            response = requests.get(url, headers=headers)
            time.sleep(1)
            if response.status_code == 200:
                response.encoding = 'utf-8'
                json_data = response.json()
                msg_list = json_data.get('msg', {})
                response_id = msg_list.get('id')
                for key, patent in unique_patents.items():
                    if patent.get('uuid') == response_id:
                        unique_patents[key].update(msg_list)
                        break
        return unique_patents
    else:
        return (f"请求失败，状态码：{response.status_code}, 响应内容：{response.text}")
# 名称：pat_name；
# 摘要：keys、zhaiyao；
# 相似度：xsd；
# 申请号：regno；
# 权限要求：zqx；
# 公开号：open_no；
# 背景技术：bg_tech；
# 技术领域：techArea；
# 附图说明：fuTuDesc；
# 权利要求书：pat_qlyqs；
# 具体实施方式：pat_examples；
# 发明内容：pat_summary；


# claimsPath， instrPath
# 调用 ollama serve 命令
# subprocess.Popen(["ollama", "serve"])
# 配置Streamlit页面
st.set_page_config(page_title="专利交底书", page_icon="📝")  # 设置界面标题和图标
st.title("📝 专利交底书")  # 显示界面标题
uploaded_file = st.file_uploader("Upload an article", type=("txt", "md", "docx"))  # 创建文件上传组件

# LLM模型选择
ollama_models = [m['name'] for m in ol.list()['models']]  # 获取所有可用的OLLAMA模型
with st.sidebar:  # 在侧边栏中显示模型选择器
    model = st.selectbox("LLM", ollama_models)  # 创建下拉菜单以选择模型

# 系统提示输入
default_prompt = (
    "你是一个专利专业人员，请阅读以上专利内容，详细回答我的任何问题，并且用中文回答我，我需要撰写一篇专利交底书，"
    "请回答以下问题，确保内容详尽，清晰，并符合专利交底书的撰写要求，根据现有内容回答，且字数尽量长，专注于数据，不要说与问题无关的话，不要自己创造问题，"
    "不要太宽泛，具体到细节领域，直接回答问题内容，不要输出其他回答，以下是我的问题：\n"
    "发明名称、技术领域：\n"
    "背景技术：\n"
    "与本发明最相近似的现有实现方案一（现有技术一的技术方案、现有技术一的缺点）：\n"
    "与本发明最相近似的现有实现方案二（现有技术二的技术方案、现有技术二的缺点）：\n"
    "发明内容（发明目的、发明方案、达到的有益效果）：\n"
    "附图说明(流程图，结构图)：\n"
    "实施方式：\n"
    "实施例（一个包含所有所知数据的流程）：\n"
    "是否还有别的替代方案同样能完成发明目的：\n"
    "本发明的技术关键点和欲保护点是什么：\n"
)
merged_prompt = re.split(r'\n', default_prompt)  # 使用正则表达式将默认提示拆分为更小的部分

system_prompt = st.sidebar.text_area("System Prompt", value=default_prompt, height=500)

# Initialize chat history
chat_key = f"对话_chat_history_{model}"
init_chat_history(chat_key, merged_prompt[0])
chat_history = st.session_state.chat_history[chat_key]

# Display chat history
for message in chat_history:
    print_chat_message(message)

# Handle file upload
if uploaded_file:
    if uploaded_file.name.endswith(".docx"):
        st.session_state.article = read_docx(uploaded_file)
    else:
        st.session_state.article = uploaded_file.read().decode()
elif 'article' not in st.session_state:
    st.session_state.article = ''

# User input for question
question = st.text_area("输入问题：", key="input0", placeholder="发明名称：\n"
                                                            "[请在此处填写发明的名称]\n" +
                                                            "技术方案或解决的问题：\n" +
                                                            "[描述发明的核心技术方案或解决的问题]\n\n\n" +
                                                            "产生的效果或解决的问题：\n" +
                                                            "[列出第一个有益效果或第一个解决的问题]\n" +
                                                            "[列出第二个有益效果或第二个解决的问题]\n" +
                                                            "[根据需要添加更多有益效果或更多解决的问题]\n" +
                                                            "具体实施例：\n" +
                                                            "[描述一个或多个具体实施例，包括关键参数、操作条件等]\n", height=200)
debug_mode = st.sidebar.checkbox("Debug Mode", value=False)

# Handle user question submission
if st.button("发送", key="button1"):
    article = st.session_state.article
    prompt = f"Here's an article:\n\n<article>\n{article}\n\n</article>\n\n{question}"

    if "article" in st.session_state:
        article = st.session_state.article

        # Ensure doc is initialized
        if "doc" not in st.session_state:
            st.session_state.doc = Document()
            st.session_state.doc.styles['Normal'].font.name = '宋体'
            st.session_state.doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            st.session_state.doc.add_picture('./sheet/SRIPPM.png', width=Inches(2.0), height=Inches(0.5))
            content = st.session_state.doc.add_heading('', level=1).add_run('专利交底书')
            content.font.name = u'宋体'
            content._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        doc = st.session_state.doc

        question = str(question)
        article = str(article)
        response_0 = search_patents(question + article)
        if isinstance(response_0, int):
            st.error(f"Error: {response_0}")
        else:
            technology_1_zhaiyao = ('专利名：' + response_0['0']['pat_name'] + '\n\n' + '专利号：' + response_0['0']['open_no'] + '\n\n' +
                            '内容：' + response_0['0']['zhaiyao'])
            technology_2_zhaiyao = ('专利名：' + response_0['1']['pat_name'] + '\n\n' + '专利号：' + response_0['1']['open_no'] + '\n\n' +
                            '内容：' + response_0['1']['zhaiyao'])

            for i in range(1, len(merged_prompt) - 1):
                if "背景技术" in merged_prompt[i]:
                    input_prompt = technology_1_zhaiyao + '\n' + technology_2_zhaiyao + '\n' + merged_prompt[0] + '\n' + merged_prompt[i]
                elif "与本发明最相近似的现有实现方案一" in merged_prompt[i]:
                    input_prompt = ("现有技术一：" + '\n\n' + technology_1_zhaiyao + '\n\n' + response_0['0']['pat_examples'] + '\n\n' +
                                    "本技术：" + '\n\n'  + question + '\n\n' + article + '\n\n'
                                    + merged_prompt[0] + '\n\n' + merged_prompt[i])
                elif "与本发明最相近似的现有实现方案二" in merged_prompt[i]:
                    input_prompt = ("现有技术二：" + '\n\n' + technology_1_zhaiyao + '\n\n' + response_0['1']['pat_examples'] + '\n\n' +
                                    "本技术：" + '\n\n'  + question + '\n\n' + article + '\n\n'
                                    + merged_prompt[0] + '\n\n' + merged_prompt[i])
                else:
                    input_prompt = question + '\n\n' + article + '\n\n' + merged_prompt[0] + '\n\n' + merged_prompt[i]

                user_message = {"role": "user", "content": input_prompt}
                print_chat_message(user_message)
                chat_history.append(user_message)

                if "背景技术" in merged_prompt[i]:
                    response = ol.chat(model=model, messages=chat_history)
                    answer = response["message"]["content"]
                    answer = (answer + '\n' + '\n' + technology_1_zhaiyao + '\n' + '\n' + technology_2_zhaiyao)
                elif "附图说明" in merged_prompt[i]:
                    answer = "流程图，结构图"
                else:
                    response = ol.chat(model=model, messages=chat_history)
                    answer = response["message"]["content"]
                run = doc.add_heading('', level=2).add_run(f'{merged_prompt[i]}')
                run.font.size = Pt(14)
                run.bold = True
                run.font.name = u'宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                doc.add_paragraph(answer)
                ai_message = {"role": "assistant", "content": answer}
                print_chat_message(ai_message)
                chat_history.append(ai_message)

            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            doc_bytes = buffer.getvalue()
            doc_bytes = remove_symbols_from_word(doc_bytes)
            st.download_button(label="下载文档", data=doc_bytes, file_name="patent_disclosure.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        # Debug info
        debug_info = {"messages": chat_history, "response": response}
        if debug_mode:
            st.write("Debug Info: Complete Prompt Interaction")
            st.json(debug_info)

        # Limit chat history
        if len(chat_history) > 30:
            chat_history = chat_history[-30:]
        st.session_state.chat_history[chat_key] = chat_history
