import streamlit as st
import ollama as ol
import datetime
import json
import requests
import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import io
from docx.shared import Inches
import subprocess

def log_interaction(action, data):
    # 日志记录功能
    timestamp = datetime.datetime.now().isoformat()  # 获取当前时间的ISO格式
    log = {"timestamp": timestamp, "action": action, "data": data}  # 创建一个包含时间戳、动作和数据的字典
    with open("user_interactions_log.json", "a") as logfile:  # 以追加模式打开日志文件
        logfile.write(json.dumps(log) + "\n")  # 将日志记录写入文件，并添加换行符


def print_txt(text):
    # 定义一个函数，以正确方向显示阿拉伯文本
    if any("\u0600" <= c <= "\u06FF" for c in text):  # 检查文本中是否包含阿拉伯字符
        text = f"<p style='direction: rtl; text-align: right;'>{text}</p>"  # 如果包含，则设置文本为从右到左显示
    st.markdown(text, unsafe_allow_html=True)  # 使用Markdown显示文本，并允许HTML


def print_chat_message(message):
    # 定义打印聊天信息的函数
    text = message["content"]  # 获取信息内容
    if message["role"] == "user":  # 如果信息来自用户
        with st.chat_message("user", avatar="🎙️"):  # 设置用户头像并显示信息
            print_txt(text)  # 调用print_txt函数以处理文本方向
    elif message["role"] == "assistant":  # 如果信息来自助手
        with st.chat_message("assistant", avatar="🦙"):  # 设置助手头像并显示信息
            print_txt(text)  # 调用print_txt函数以处理文本方向


def init_chat_history(key, system_prompt):
    # 定义打印聊天信息的函数
    if "chat_history" not in st.session_state:  # 如果会话状态中没有聊天记录
        st.session_state.chat_history = {}  # 初始化聊天记录
    if key not in st.session_state.chat_history:  # 如果指定的聊天记录不存在
        st.session_state.chat_history[key] = [{"role": "system", "content": system_prompt}]  # 创建新的聊天记录并添加系统提示


def read_docx(file):
    # 读取DOCX文件
    doc = Document(file)  # 打开Word文档
    full_text = []  # 初始化一个空列表以存储文本
    for para in doc.paragraphs:  # 遍历文档中的所有段落
        full_text.append(para.text)  # 将段落文本添加到列表中
    return '\n'.join(full_text)  # 将所有段落文本合并为一个字符串并返回

def search_patents(x1):
    url = "https://open.cnipr.com/cnipr-api/v1/api/search/sf1/BC816A216BF85BBA27E7D1165EEFA953"

    # 构建请求数据
    data = {
        "client_id": "BC816A216BF85BBA27E7D1165EEFA953",
        "openid": "621A46A70CD30A72D7AB3F806B32C823",
        "access_token": "EAA108A45D99233AC4855E75B9C2FBCB",
        "exp": f"ss = (’{x1}‘)",
        "dbs": ["FMZL"],
        "from": 0,
        "size": 2,
        "option": 2,
        "highLight": False,
        "isDbAgg": False,
        "displayCols": "title,abs,claimsPath"
    }

    # 设置请求头
    headers = {
        "Content-Type": "application/x-www-form-urlencoded"
    }
    # 发送POST请求
    response = requests.post(url, data=data, headers=headers)
    # 打印响应结果
    if response.status_code == 200:
        return response.json()
    else:
        return (f"请求失败，状态码：{response.status_code}, 响应内容：{response.text}")

# claimsPath， instrPath
# 调用 ollama serve 命令
# subprocess.Popen(["ollama", "serve"])
# 配置Streamlit页面
st.set_page_config(page_title="专利交底书", page_icon="📝")  # 设置界面标题和图标
st.title("📝 专利交底书")  # 显示界面标题
uploaded_file = st.file_uploader("Upload an article", type=("txt", "md", "docx"))  # 创建文件上传组件

# LLM模型选择
ollama_models = [m['name'] for m in ol.list()['models']]  # 获取所有可用的OLLAMA模型
with st.sidebar:  # 在侧边栏中显示模型选择器
    model = st.selectbox("LLM", ollama_models)  # 创建下拉菜单以选择模型

# 系统提示输入
default_prompt = (
    "你是一名资深软件测试工程师，"
    "具有丰富的经验和深入的专业知识。你的主要职责和特点包括，"
    "深刻理解软件需求，你能够准确理解和分析软件的功能需求。你知道如何将复杂的需求分解为具体的测试点，并能够识别潜在的风险和挑战。"
    "制定测试策略，你擅长规划和设计测试策略。这包括选择适当的测试类型（如功能测试，接口测试等），确定测试范围和目标，以及选择最合适的测试方法和工具。"
    "细化测试项，你能够详细定义测试项，包括为每个测试项命名和标识，制定具体的测试计划。你知道如何设定有效的测试标准，以确保软件质量。"
    "确保实用性和可执行性，你生成的测试大纲都是实用和可执行的。你的目标是确保测试团队能够轻松理解和实施这些策略，并通过这些测试确保软件的质量和性能。"
    "任务描述，"
    "基于上述角色描述，接下来你将会接收软件功能需求的描述。你需要按照以下分步骤的方法，对用户的输入进行分析，"
    "识别和理解软件需求，理解软件需要实现的功能和要求。"
    "将需求分解为具体的测试点，将功能需求细分为具体的测试点。"
    "识别潜在的风险和挑战，识别可能影响软件质量的风险和挑战。"
    "规划和设计测试策略，设计合适的测试策略以覆盖所有功能需求。"
    "选择适当的测试类型和方法，选择最适合的软件测试类型和方法。"
    "定义测试项并制定测试计划，详细定义每个测试项并制定具体的测试计划。"
    "设定测试标准以确保软件质量，设定明确的测试标准，确保软件功能和性能的可靠性。"
    "基于以上数据，请回答我的问题。以下是我的问题："
    "测评范围、\n"
    "测试类型、\n"
    "功能测试的唯一标识，"
    "测试项一（测试项名称，测试项标识，，优先级，追踪关系，需求描述，测试项描述，测试方法，充分性要求，通过准则），"
    "测试项二（测试项名称，测试项标识，，优先级，追踪关系，需求描述，测试项描述，测试方法，充分性要求，通过准则），"
    "测试项三（测试项名称，测试项标识，，优先级，追踪关系，需求描述，测试项描述，测试方法，充分性要求，通过准则），"
    "测试项四（测试项名称，测试项标识，，优先级，追踪关系，需求描述，测试项描述，测试方法，充分性要求，通过准则），"
    "测试项五（测试项名称，测试项标识，，优先级，追踪关系，需求描述，测试项描述，测试方法，充分性要求，通过准则），"
    "描述，测试步骤，风险管理、\n"
    "性能测试的唯一标识，"
    "测试项一（测试项名称，测试项标识，，优先级，追踪关系，需求描述，测试项描述，测试方法，充分性要求，通过准则），"
    "测试项二（测试项名称，测试项标识，，优先级，追踪关系，需求描述，测试项描述，测试方法，充分性要求，通过准则），"
    "测试项三（测试项名称，测试项标识，，优先级，追踪关系，需求描述，测试项描述，测试方法，充分性要求，通过准则），"
    "测试项四（测试项名称，测试项标识，，优先级，追踪关系，需求描述，测试项描述，测试方法，充分性要求，通过准则），"
    "测试项五（测试项名称，测试项标识，，优先级，追踪关系，需求描述，测试项描述，测试方法，充分性要求，通过准则），"
    "描述，测试步骤，风险管理、\n"
    "接口测试的唯一标识，"
    "测试项一（测试项名称，测试项标识，，优先级，追踪关系，需求描述，测试项描述，测试方法，充分性要求，通过准则），"
    "测试项二（测试项名称，测试项标识，，优先级，追踪关系，需求描述，测试项描述，测试方法，充分性要求，通过准则），"
    "测试项三（测试项名称，测试项标识，，优先级，追踪关系，需求描述，测试项描述，测试方法，充分性要求，通过准则），"
    "测试项四（测试项名称，测试项标识，，优先级，追踪关系，需求描述，测试项描述，测试方法，充分性要求，通过准则），"
    "测试项五（测试项名称，测试项标识，，优先级，追踪关系，需求描述，测试项描述，测试方法，充分性要求，通过准则），"
    "描述，测试步骤，风险管理、\n"
    "问题严重性等级、\n"
    "问题处理方法、\n"
)
split_prompt = re.split(r'(：|、)', default_prompt)  # 使用正则表达式将默认提示拆分为更小的部分
merged_prompt = []
for i in range(0, len(split_prompt) - 1, 2):  #
    merged_prompt.append(split_prompt[i] + split_prompt[i + 1])
if len(split_prompt) % 2 != 0:
    merged_prompt.append(split_prompt[-1])

system_prompt = st.sidebar.text_area("System Prompt", value=default_prompt, height=500)

# Initialize chat history
chat_key = f"对话_chat_history_{model}"
init_chat_history(chat_key, merged_prompt[0])
chat_history = st.session_state.chat_history[chat_key]

# Display chat history
for message in chat_history:
    print_chat_message(message)

# Handle file upload
if uploaded_file:
    if uploaded_file.name.endswith(".docx"):
        st.session_state.article = read_docx(uploaded_file)
    else:
        st.session_state.article = uploaded_file.read().decode()
elif 'article' not in st.session_state:
    st.session_state.article = ''

# User input for question
question = st.text_area("输入问题：", key="input0", placeholder="名称：（请填写方案的名称）\n"
                                                               "方案：（简要描述方案的原理和整体流程）\n"
                                                               "步骤1：预处理（具体描述预处理的操作，包括数据清洗、格式转换等）\n"
                                                               "步骤2：（描述第二步的操作，包含具体的技术或方法）\n"
                                                               "步骤3：（继续描述后续步骤，依此类推）\n"
                                                               "效果：\n"
                                                               "预处理效果：（描述预处理后数据的改善或变化，例如噪声减少、数据完整性提升等）\n"
                                                               "步骤2效果：（描述第二步的具体效果，例如模型训练效果提升、准确率变化等）\n"
                                                               "步骤3效果：（继续描述每个步骤的具体效果，依此类推）\n"
                                                               "数据：\n"
                                                               "步骤1开始到结束的数据：时间： （填写具体的时间范围或处理时间）温度： （填写实验或处理过程中涉及的温度范围）浓度： （填写相关物质的浓度范围）\n"
                                                               "步骤2的数据：时间： （填写具体的时间范围或处理时间）温度： （填写实验或处理过程中涉及的温度范围）浓度： （填写相关物质的浓度范围）\n"
                                                               "步骤3的数据：时间： （填写具体的时间范围或处理时间）温度： （填写实验或处理过程中涉及的温度范围）浓度： （填写相关物质的浓度范围）\n", height=200)
debug_mode = st.sidebar.checkbox("Debug Mode", value=False)

# Handle user question submission
if st.button("发送", key="button1"):
    article = st.session_state.article
    prompt = f"Here's an article:\n\n<article>\n{article}\n\n</article>\n\n{question}"

    if "article" in st.session_state:
        article = st.session_state.article

        # Ensure doc is initialized
        if "doc" not in st.session_state:
            st.session_state.doc = Document()
            st.session_state.doc.styles['Normal'].font.name = '宋体'
            st.session_state.doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            st.session_state.doc.add_picture('./sheet/SRIPPM.png', width=Inches(2.0), height=Inches(0.5))
            content = st.session_state.doc.add_heading('', level=1).add_run('专利交底书')
            content.font.name = u'宋体'
            content._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        doc = st.session_state.doc

        question = str(question)
        article = str(article)
        for i in range(1, len(merged_prompt) - 1):
            input_prompt = question + '\n' + article + '\n' + merged_prompt[0] + '\n' + merged_prompt[i]

            user_message = {"role": "user", "content": input_prompt}
            print_chat_message(user_message)
            chat_history.append(user_message)

            response = ol.chat(model=model, messages=chat_history)
            answer = response["message"]["content"]

            run = doc.add_heading('', level=2).add_run(f'{merged_prompt[i]}')
            run.font.size = Pt(14)
            run.bold = True
            run.font.name = u'宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            doc.add_paragraph(answer)
            ai_message = {"role": "assistant", "content": answer}
            print_chat_message(ai_message)
            chat_history.append(ai_message)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        doc_bytes = buffer.getvalue()
        st.download_button(label="下载文档", data=doc_bytes, file_name="patent_disclosure.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        # Debug info
        debug_info = {"messages": chat_history, "response": response}
        if debug_mode:
            st.write("Debug Info: Complete Prompt Interaction")
            st.json(debug_info)

        # Limit chat history
        if len(chat_history) > 30:
            chat_history = chat_history[-30:]
        st.session_state.chat_history[chat_key] = chat_history
